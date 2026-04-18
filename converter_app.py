"""
MD Converter — macOS App
DOCX / PDF / HWPX / Text → Markdown
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import os
import re
import zipfile
import xml.etree.ElementTree as ET

# ── 변환 로직 ──────────────────────────────────────────────

def docx_to_md(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx가 설치되어 있지 않습니다.")

    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        style = para.style.name or ""
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        if style.startswith("Heading 1") or style == "제목 1":
            lines.append(f"# {text}")
        elif style.startswith("Heading 2") or style == "제목 2":
            lines.append(f"## {text}")
        elif style.startswith("Heading 3") or style == "제목 3":
            lines.append(f"### {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            md_text = ""
            for run in para.runs:
                t = run.text
                if run.bold and run.italic:
                    t = f"***{t}***"
                elif run.bold:
                    t = f"**{t}**"
                elif run.italic:
                    t = f"*{t}*"
                md_text += t
            lines.append(md_text if md_text.strip() else text)

    for table in doc.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("")
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines)


def pdf_to_md(path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber가 설치되어 있지 않습니다.")

    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
            for table in page.extract_tables():
                if not table:
                    continue
                header = [str(c or "").strip() for c in table[0]]
                lines.append("")
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                for row in table[1:]:
                    cells = [str(c or "").strip() for c in row]
                    lines.append("| " + " | ".join(cells) + " |")
                lines.append("")

    return "\n".join(lines)


def hwpx_to_md(path: str) -> str:
    """HWPX는 ZIP 컨테이너 안에 XML. 스타일/개요(outline) 정보로 제목 레벨 판단."""
    paragraphs = []  # [(level, text)] — level: 0=본문, 1~6=제목

    try:
        with zipfile.ZipFile(path, 'r') as z:
            # 본문 XML들
            xml_files = sorted([n for n in z.namelist()
                                if n.startswith("Contents/section") and n.endswith(".xml")])
            if not xml_files:
                xml_files = [n for n in z.namelist() if n.endswith(".xml")]

            # 헤더(스타일 정의) 읽기 — 제목 스타일 ID 매핑
            style_to_level = {}
            header_files = [n for n in z.namelist() if "header" in n.lower() and n.endswith(".xml")]
            for hf in header_files:
                try:
                    with z.open(hf) as f:
                        header_xml = f.read().decode("utf-8", errors="ignore")
                    hroot = ET.fromstring(header_xml)
                    for elem in hroot.iter():
                        tag = elem.tag.split("}")[-1]
                        if tag == "style":
                            sid = elem.attrib.get("id", "")
                            name = elem.attrib.get("name", "").lower()
                            # 한글/영문 제목 스타일 인식
                            if "개요 1" in name or "heading 1" in name or name == "제목 1":
                                style_to_level[sid] = 1
                            elif "개요 2" in name or "heading 2" in name or name == "제목 2":
                                style_to_level[sid] = 2
                            elif "개요 3" in name or "heading 3" in name or name == "제목 3":
                                style_to_level[sid] = 3
                            elif "개요 4" in name or "heading 4" in name:
                                style_to_level[sid] = 4
                            elif "개요 5" in name or "heading 5" in name:
                                style_to_level[sid] = 5
                            elif name == "제목":
                                style_to_level[sid] = 1
                except ET.ParseError:
                    continue

            # 본문 파싱
            for xml_name in xml_files:
                try:
                    with z.open(xml_name) as f:
                        content = f.read().decode("utf-8", errors="ignore")
                    root = ET.fromstring(content)

                    for elem in root.iter():
                        tag = elem.tag.split("}")[-1]
                        if tag == "p":  # 문단
                            style_id = elem.attrib.get("paraPrIDRef", "")
                            level = style_to_level.get(style_id, 0)

                            # 문단 안의 모든 텍스트 합치기
                            para_text = ""
                            for sub in elem.iter():
                                sub_tag = sub.tag.split("}")[-1]
                                if sub_tag == "t" and sub.text:
                                    para_text += sub.text

                            para_text = para_text.strip()
                            if para_text:
                                paragraphs.append((level, para_text))
                            else:
                                paragraphs.append((0, ""))  # 빈 줄
                except ET.ParseError:
                    continue
    except zipfile.BadZipFile:
        raise ValueError("올바른 HWPX 파일이 아닙니다.")

    # 마크다운 조립
    lines = []
    for level, text in paragraphs:
        if not text:
            lines.append("")
            continue
        if level >= 1 and level <= 6:
            lines.append(f"{'#' * level} {text}")
        else:
            # 스타일 정보 없는 경우에만 휴리스틱 적용 (최후 수단)
            lines.append(text)

    return "\n".join(lines)


def text_to_md(text: str) -> str:
    lines = text.splitlines()
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append("")
            continue
        if len(stripped) < 40 and not stripped.endswith((".", ",", "?", "!")):
            result.append(f"## {stripped}")
        else:
            result.append(stripped)
    return "\n".join(result)


def clean_markdown(md: str) -> str:
    md = re.sub(r'\n{3,}', '\n\n', md)
    return md.strip()


# ── UI ─────────────────────────────────────────────────────

DARK_BG   = "#1a1a1a"
PANEL_BG  = "#242424"
BORDER    = "#333333"
ACCENT    = "#7eceff"
TEXT_MAIN = "#e8e8e8"
TEXT_DIM  = "#888888"
GREEN     = "#5ddb8e"


class MDConverterApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("MD Converter")
        self.geometry("920x680")
        self.minsize(760, 560)
        self.configure(bg=DARK_BG)
        self._result_md = ""
        self._txt_result_md = ""
        self._build_ui()
        self._check_deps()

    def _build_ui(self):
        hdr = tk.Frame(self, bg=DARK_BG, pady=16)
        hdr.pack(fill="x", padx=28)
        tk.Label(hdr, text="MD Converter", font=("Helvetica Neue", 22, "bold"),
                 fg=ACCENT, bg=DARK_BG).pack(side="left")
        tk.Label(hdr, text="DOCX · PDF · HWPX · Text  →  Markdown",
                 font=("Helvetica Neue", 12), fg=TEXT_DIM, bg=DARK_BG).pack(side="left", padx=14, pady=4)

        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True)

        self._style_notebook()

        self.tab_file = tk.Frame(nb, bg=DARK_BG)
        self.tab_text = tk.Frame(nb, bg=DARK_BG)
        nb.add(self.tab_file, text="  파일 변환  ")
        nb.add(self.tab_text, text="  텍스트 붙여넣기  ")

        self._build_file_tab(self.tab_file)
        self._build_text_tab(self.tab_text)

        # 하단 저작권 표시
        footer = tk.Frame(self, bg=DARK_BG)
        footer.pack(side="bottom", fill="x", pady=(0, 8))
        tk.Label(
            footer,
            text="© Renai. All rights reserved.  ·  Moonlight Archive",
            font=("Helvetica Neue", 10),
            fg=TEXT_DIM, bg=DARK_BG
        ).pack()

    def _style_notebook(self):
        s = ttk.Style(self)
        s.theme_use("default")
        s.configure("TNotebook", background=DARK_BG, borderwidth=0, tabmargins=0)
        s.configure("TNotebook.Tab",
                    background=PANEL_BG, foreground=TEXT_DIM,
                    font=("Helvetica Neue", 12), padding=(16, 8),
                    borderwidth=0)
        s.map("TNotebook.Tab",
              background=[("selected", DARK_BG)],
              foreground=[("selected", ACCENT)])

    def _build_file_tab(self, parent):
        drop_frame = tk.Frame(parent, bg=PANEL_BG, bd=0, relief="flat",
                              highlightbackground=BORDER, highlightthickness=1)
        drop_frame.pack(fill="x", padx=28, pady=(20, 0))

        self.drop_label = tk.Label(
            drop_frame,
            text="📂  파일을 클릭하여 선택\n\n.docx   .pdf   .hwpx",
            font=("Helvetica Neue", 14), fg=TEXT_DIM, bg=PANEL_BG,
            pady=32, cursor="hand2"
        )
        self.drop_label.pack(fill="x")
        self.drop_label.bind("<Button-1>", lambda e: self._pick_file())

        self.file_info = tk.Label(parent, text="", font=("Menlo", 11),
                                  fg=ACCENT, bg=DARK_BG)
        self.file_info.pack(pady=(8, 0))

        self.btn_convert = self._btn(parent, "변환하기", self._convert_file)
        self.btn_convert.pack(pady=(12, 0))

        self._build_result_area(parent)

    def _build_text_tab(self, parent):
        tk.Label(parent, text="텍스트를 붙여넣으세요",
                 font=("Helvetica Neue", 12), fg=TEXT_DIM, bg=DARK_BG).pack(
            anchor="w", padx=28, pady=(20, 6))

        self.text_input = tk.Text(
            parent, height=10, bg=PANEL_BG, fg=TEXT_MAIN,
            font=("Menlo", 12), insertbackground=ACCENT,
            relief="flat", bd=0, padx=12, pady=10,
            highlightbackground=BORDER, highlightthickness=1
        )
        self.text_input.pack(fill="x", padx=28)

        self._btn(parent, "변환하기", self._convert_text).pack(pady=(12, 0))

        self._build_result_area(parent, prefix="txt")

    def _build_result_area(self, parent, prefix="file"):
        tk.Frame(parent, bg=BORDER, height=1).pack(fill="x", padx=28, pady=(20, 0))

        label_row = tk.Frame(parent, bg=DARK_BG)
        label_row.pack(fill="x", padx=28, pady=(8, 4))
        tk.Label(label_row, text="Markdown 결과", font=("Helvetica Neue", 11),
                 fg=TEXT_DIM, bg=DARK_BG).pack(side="left")

        btn_row = tk.Frame(parent, bg=DARK_BG)
        btn_row.pack(fill="x", padx=28)
        self._small_btn(btn_row, "클립보드 복사", lambda: self._copy(prefix)).pack(side="left", padx=(0, 8))
        self._small_btn(btn_row, ".md 저장", lambda: self._save(prefix)).pack(side="left")

        box = tk.Text(
            parent, height=12, bg=PANEL_BG, fg=TEXT_MAIN,
            font=("Menlo", 11), insertbackground=ACCENT,
            relief="flat", bd=0, padx=12, pady=10,
            highlightbackground=BORDER, highlightthickness=1,
            state="disabled"
        )
        box.pack(fill="both", expand=True, padx=28, pady=(8, 20))

        if prefix == "txt":
            self._txt_result_box = box
        else:
            self._file_result_box = box

    def _btn(self, parent, text, cmd):
        return tk.Button(
            parent, text=text, command=cmd,
            bg=ACCENT, fg="#0a0a0a", font=("Helvetica Neue", 13, "bold"),
            relief="flat", bd=0, padx=28, pady=10, cursor="hand2",
            activebackground="#a8e6ff", activeforeground="#0a0a0a"
        )

    def _small_btn(self, parent, text, cmd):
        return tk.Button(
            parent, text=text, command=cmd,
            bg=BORDER, fg=TEXT_MAIN, font=("Helvetica Neue", 11),
            relief="flat", bd=0, padx=12, pady=5, cursor="hand2",
            activebackground="#444", activeforeground=TEXT_MAIN
        )

    def _pick_file(self):
        path = filedialog.askopenfilename(
            filetypes=[
                ("지원 파일", "*.docx *.pdf *.hwpx"),
                ("DOCX", "*.docx"),
                ("PDF", "*.pdf"),
                ("HWPX", "*.hwpx"),
            ]
        )
        if path:
            self._selected_file = path
            self.file_info.config(text=f"✓  {os.path.basename(path)}", fg=GREEN)

    def _convert_file(self):
        path = getattr(self, "_selected_file", None)
        if not path:
            messagebox.showwarning("파일 없음", "파일을 먼저 선택하세요.")
            return
        self._run_in_thread(self._do_convert_file, path)

    def _do_convert_file(self, path):
        try:
            ext = os.path.splitext(path)[1].lower()
            if ext == ".docx":
                md = docx_to_md(path)
            elif ext == ".pdf":
                md = pdf_to_md(path)
            elif ext == ".hwpx":
                md = hwpx_to_md(path)
            else:
                raise ValueError(f"지원하지 않는 형식: {ext}")
            md = clean_markdown(md)
            self._result_md = md
            self._set_result(self._file_result_box, md)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("변환 오류", str(e)))

    def _convert_text(self):
        text = self.text_input.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning("입력 없음", "텍스트를 입력하세요.")
            return
        self._run_in_thread(self._do_convert_text, text)

    def _do_convert_text(self, text):
        try:
            md = text_to_md(text)
            md = clean_markdown(md)
            self._txt_result_md = md
            self._set_result(self._txt_result_box, md)
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("변환 오류", str(e)))

    def _set_result(self, box, md):
        def _update():
            box.config(state="normal")
            box.delete("1.0", "end")
            box.insert("1.0", md)
            box.config(state="disabled")
        self.after(0, _update)

    def _copy(self, prefix):
        md = self._txt_result_md if prefix == "txt" else self._result_md
        if not md:
            messagebox.showinfo("없음", "변환 결과가 없습니다.")
            return
        self.clipboard_clear()
        self.clipboard_append(md)
        messagebox.showinfo("복사 완료", "클립보드에 복사되었습니다.")

    def _save(self, prefix):
        md = self._txt_result_md if prefix == "txt" else self._result_md
        if not md:
            messagebox.showinfo("없음", "변환 결과가 없습니다.")
            return
        save_path = filedialog.asksaveasfilename(
            defaultextension=".md",
            filetypes=[("Markdown", "*.md"), ("Text", "*.txt")]
        )
        if save_path:
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(md)
            messagebox.showinfo("저장 완료", f"저장되었습니다:\n{save_path}")

    def _run_in_thread(self, fn, *args):
        t = threading.Thread(target=fn, args=args, daemon=True)
        t.start()

    def _check_deps(self):
        missing = []
        try:
            import docx
        except ImportError:
            missing.append("python-docx")
        try:
            import pdfplumber
        except ImportError:
            missing.append("pdfplumber")
        if missing:
            msg = "다음 패키지가 없습니다:\n" + "\n".join(f"  pip install {p}" for p in missing)
            self.after(500, lambda: messagebox.showwarning("패키지 누락", msg))


if __name__ == "__main__":
    app = MDConverterApp()
    app.mainloop()
